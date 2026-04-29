from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
import os

def create_massive_document():
    doc = Document()
    
    # Configure global styling for University Standard (12pt font, 1.5 line spacing)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Helper functions
    def add_heading(text, level, color=(0, 51, 102), page_break_before=False):
        if page_break_before:
            doc.add_page_break()
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Arial'
            if color:
                run.font.color.rgb = RGBColor(*color)
            if level == 1:
                run.bold = True
                run.font.size = Pt(16)
        return h

    def add_blank_page(title):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for _ in range(8): p.add_run("\n")
        run = p.add_run(f"[ THIS PAGE IS RESERVED FOR: ]\n\n{title}")
        run.bold = True
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(127, 140, 141)

    # --- CHAPTER 1: INITIAL PAGES (4 Pages) ---
    add_blank_page("(1) Project Title Page")
    add_blank_page("(2) College Originality Certificate")
    add_blank_page("(3) Internship Allotment letter")
    add_blank_page("(4) Internship Completion Certificate")

    # --- PREFACE & ACKNOWLEDGEMENT (2 Pages) ---
    add_heading("Preface", level=1, page_break_before=True)
    doc.add_paragraph("The rapid advancement of Information Technology has profoundly transformed various sectors, with the educational domain being one of the most significantly impacted. This project report on 'Gurukul Edu WebApp' encapsulates the journey of developing a robust, scalable, and highly efficient educational management system. The fundamental premise of this project is to bridge the technological gap in traditional academic administration by introducing a centralized digital platform.")
    doc.add_paragraph("Throughout the development lifecycle, modern software engineering paradigms were rigorously followed. The project leverages the MERN stack—MongoDB, Express.js, React.js, and Node.js—to deliver a real-time, responsive, and secure user experience. This document serves as a comprehensive guide, detailing the entire Software Development Life Cycle (SDLC) from feasibility analysis and requirement gathering to system design, implementation, and rigorous software testing.")

    add_heading("Acknowledgement", level=1, page_break_before=True)
    doc.add_paragraph("The successful realization of this extensive project is the culmination of immense support, guidance, and encouragement from numerous individuals and institutions. I take this opportunity to express my profound gratitude to everyone who contributed to this endeavor.")
    doc.add_paragraph("First and foremost, I extend my deepest appreciation to the esteemed management and faculty members of my college. Their dedication to academic excellence provided me with the theoretical foundation and the analytical mindset required to execute complex software architectures.")
    doc.add_paragraph("I am deeply indebted to my Project Guide for their continuous mentorship. Their constructive feedback, technical insights into full-stack development, and strategic direction were the true cornerstones of this project's success. Their willingness to assist during critical debugging phases was invaluable.")
    doc.add_paragraph("Lastly, I would like to thank my peers, friends, and family for their unwavering moral support. Developing an enterprise-grade application requires immense patience and perseverance, and their encouragement kept me motivated throughout this challenging yet highly rewarding journey.")

    # --- INDEX (2 Pages) ---
    add_heading("Table of Contents", level=1, page_break_before=True)
    index_items = [
        "1. Introduction & Abstract",
        "2. Motivation & Problem Statement",
        "3. System Scope & Limitations",
        "4. System Analysis & Feasibility Study",
        "   4.1 Technical Feasibility",
        "   4.2 Economic Feasibility",
        "   4.3 Operational Feasibility",
        "   4.4 Schedule Feasibility",
        "5. Hardware & Software Specifications",
        "   5.1 Technology Stack Justification",
        "6. System Design (UML Diagrams)",
        "   6.1 System Flowchart",
        "   6.2 Data Flow Diagrams (DFD 0, 1, 2)",
        "   6.3 Entity-Relationship (ER) Diagram",
        "   6.4 UML Class & Sequence Diagrams",
        "   6.5 UML Activity Diagram",
        "   6.6 Detailed Use Case Descriptions",
        "7. Database Design & Data Dictionary",
        "8. Implementation & MVC Architecture",
        "9. Software Testing & Test Cases",
        "10. System Screenshots",
        "11. Conclusion & Future Enhancements",
        "12. Bibliography"
    ]
    for item in index_items:
        doc.add_paragraph(item, style='List Number')

    # --- CHAPTER 2: INTRO (3 Pages) ---
    add_heading("1. Introduction & Abstract", level=1, page_break_before=True)
    doc.add_paragraph("The Gurukul Edu WebApp is a comprehensive, cloud-based Educational Resource Planning (ERP) system designed to digitalize and automate the day-to-day academic and administrative workflows of educational institutions. In an era where data velocity is critical, traditional paper-based systems or isolated Excel spreadsheets are no longer sufficient. This system provides a unified interface for Administrators, Teachers, and Students.")
    doc.add_paragraph("Abstract: The system is built using a decoupled architecture where the React.js frontend communicates with a Node.js/Express RESTful API, supported by a MongoDB NoSQL database. Core modules include Role-Based Access Control (RBAC), Attendance Tracking, Study Material Distribution via Cloudinary CDN, and Financial (Fee) Management.")

    add_heading("2. Motivation & Problem Statement", level=1, page_break_before=True)
    doc.add_paragraph("Motivation: The primary motivation behind developing this system was observing the massive operational inefficiencies in local educational institutes. Teachers spend an excessive amount of time manually taking attendance and distributing physical notes, which eats into instructional time. Administrators struggle to track fee defaulters, leading to financial discrepancies.")
    doc.add_paragraph("Problem Statement: 'To design and implement a centralized web application that eliminates manual record-keeping, ensures secure role-based data access, automates academic resource distribution, and provides real-time analytical dashboards for institutional management.'")

    add_heading("3. System Scope & Limitations", level=1, page_break_before=True)
    doc.add_paragraph("Scope:\n- Complete automation of user onboarding and batch assignments.\n- Real-time digital attendance marking and tracking.\n- Secure upload and download of academic materials (PDFs, Notes) using cloud storage.\n- Centralized fee tracking system for administrative oversight.")
    doc.add_paragraph("Limitations:\n- The system currently requires an active internet connection as it lacks a robust offline synchronization mechanism.\n- Direct online payment gateways (like Razorpay) are not integrated in the current version; fee statuses are manually updated by the admin.\n- The system is optimized for web browsers and does not yet have a dedicated native mobile application (Android/iOS).")

    # --- CHAPTER 3: FEASIBILITY (3 Pages) ---
    add_heading("4. System Analysis & Feasibility Study", level=1, page_break_before=True)
    doc.add_paragraph("Before initiating the development phase, an exhaustive feasibility study was conducted to determine if the proposed system was viable, cost-effective, and technically possible within the given constraints.")

    add_heading("4.1 Technical Feasibility", level=2)
    doc.add_paragraph("Technical feasibility evaluates the current technical resources and technology stack required. The MERN stack is entirely open-source and possesses massive community support. Deployment platforms such as Render (for backend backend), Vercel (for frontend), and MongoDB Atlas (for database) provide free tiers that are more than capable of handling the prototype phase. Thus, the project is highly technically feasible.")

    add_heading("4.2 Economic Feasibility", level=2, page_break_before=True)
    doc.add_paragraph("Economic feasibility involves a cost-benefit analysis. Since the entire technology stack (Node.js, React.js, MongoDB) consists of free, open-source software, there are zero licensing costs. The only recurring costs would be premium cloud hosting if the application scales to thousands of concurrent users. The benefits (time saved, reduced paper costs, error reduction) heavily outweigh the development and hosting costs, rendering it economically highly viable.")

    add_heading("4.3 Operational Feasibility", level=2)
    doc.add_paragraph("Operational feasibility measures how well the proposed system solves the problems and how easily it can be adopted by the end-users. The React frontend is designed following modern UX/UI guidelines using Tailwind CSS, making it extremely intuitive. Teachers and admins require less than an hour of training to fully understand the dashboard workflows. Therefore, user resistance will be minimal.")

    add_heading("4.4 Schedule Feasibility", level=2)
    doc.add_paragraph("The project was planned using Agile methodologies. Divided into sprints, the core authentication module took 1 week, database schema design 1 week, API development 2 weeks, and frontend integration 2 weeks. The timeline was strictly adhered to, proving high schedule feasibility.")

    # --- CHAPTER 4: HW/SW SPECS (3 Pages) ---
    add_heading("5. Hardware & Software Specifications", level=1, page_break_before=True)
    
    add_heading("Client-Side Requirements", level=2)
    doc.add_paragraph("- Hardware: Any standard PC, Laptop, Tablet, or Smartphone with at least 2GB RAM.\n- Software: Modern Web Browsers supporting HTML5 and ES6 JavaScript (Google Chrome, Mozilla Firefox, Safari, Edge).\n- Network: Standard 3G/4G or broadband internet connection (Min 1 Mbps).")

    add_heading("Server-Side Requirements", level=2)
    doc.add_paragraph("- Hardware: Cloud-hosted Virtual Private Server (VPS) with minimum 1 vCPU and 512MB RAM.\n- Software Environment: Node.js Runtime (v18.x or higher).\n- Database: MongoDB Atlas Cluster (M0 Sandbox minimum).\n- Storage: Cloudinary API for secure media hosting.")

    add_heading("5.1 Technology Stack Justification", level=2, page_break_before=True)
    doc.add_paragraph("1. MongoDB (Database): Unlike SQL databases, MongoDB's document-oriented architecture allows for flexible schema design. Educational data is often unstructured or hierarchical (e.g., varying test formats, dynamic arrays of enrolled students), making NoSQL the perfect fit.")
    doc.add_paragraph("2. Express.js & Node.js (Backend): Node.js uses an event-driven, non-blocking I/O model which is exceptionally lightweight and efficient for data-intensive real-time applications. Express simplifies routing and middleware integration (like JWT validation).")
    doc.add_paragraph("3. React.js (Frontend): React's Virtual DOM provides unparalleled performance by updating only the components that change, rather than reloading the entire page. This results in a seamless, app-like experience for the user.")
    doc.add_paragraph("4. Tailwind CSS (Styling): Utility-first CSS framework that drastically speeds up UI development. It ensures the application is 100% responsive across all device sizes without writing thousands of lines of custom CSS media queries.")

    # --- CHAPTER 5: DESIGN & DIAGRAMS (10+ Pages) ---
    def add_diagram_placeholder(title):
        add_heading(title, level=1, page_break_before=True)
        p = doc.add_paragraph(f"\n\n\n<< PLEASE PASTE THE '{title}' IMAGE HERE FROM diagrams_massive.html >>\n\n\n")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold, p.runs[0].font.color.rgb = True, RGBColor(231, 76, 60)

    add_diagram_placeholder("6.1 System Flowchart")
    add_diagram_placeholder("6.2 Data Flow Diagram (Level 0)")
    add_diagram_placeholder("6.2 Data Flow Diagram (Level 1)")
    add_diagram_placeholder("6.2 Data Flow Diagram (Level 2)")
    add_diagram_placeholder("6.3 Entity-Relationship (ER) Diagram")
    add_diagram_placeholder("6.4 UML Class Diagram")
    add_diagram_placeholder("6.4 UML Sequence Diagram")
    add_diagram_placeholder("6.5 UML Activity Diagram")

    # Use Case Descriptions (Takes up 3 pages easily)
    add_heading("6.6 Detailed Use Case Descriptions", level=1, page_break_before=True)
    doc.add_paragraph("A Use Case Description is a formal, textual representation of a specific interaction between the actor and the system. Below are the exhaustive descriptions for core system modules.")

    def add_uc_table(uc_id, name, actor, pre, flow, post):
        add_heading(f"Use Case: {uc_id} - {name}", level=3)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'
        data = [("Primary Actor", actor), ("Pre-conditions", pre), ("Main Success Flow", flow), ("Post-conditions", post)]
        for i, (k, v) in enumerate(data):
            table.rows[i].cells[0].text = k
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            table.rows[i].cells[1].text = v
        doc.add_paragraph("\n")

    add_uc_table("UC-01", "System Authentication (Login)", "Admin, Teacher, Student", "User must be registered in the database.", "1. User navigates to login page.\n2. User enters Email and Password.\n3. System hashes password and compares with DB.\n4. System generates JWT and sends it to client.\n5. System redirects user to respective dashboard.", "User session is created. JWT is stored in local storage.")
    doc.add_page_break()
    add_uc_table("UC-02", "Upload Academic Material", "Teacher", "Teacher must be logged in and assigned to a batch.", "1. Teacher navigates to Academic Hub.\n2. Teacher selects Batch, enters Title, and selects File (PDF/Image).\n3. System uploads file to Cloudinary API.\n4. Cloudinary returns secure HTTPS URL.\n5. System saves URL and metadata to MongoDB.\n6. System shows success message.", "Material record is successfully inserted into the database and visible to students.")
    add_uc_table("UC-03", "Mark Attendance", "Teacher", "Teacher must be logged in.", "1. Teacher selects a specific Batch and Date.\n2. System fetches list of all enrolled students.\n3. Teacher marks status (Present/Absent) for each student.\n4. Teacher clicks Submit.\n5. System bulk-inserts attendance records into DB.", "Database is updated with new attendance documents for the specified date.")
    doc.add_page_break()
    add_uc_table("UC-04", "View Fee Status", "Student", "Student must be logged in.", "1. Student navigates to Finance Dashboard.\n2. System extracts Student ID from JWT.\n3. System queries Fee collection for records matching ID.\n4. System calculates Total Amount minus Amount Paid.\n5. System renders Due Amount on UI.", "Student successfully views their financial status.")

    # --- CHAPTER 6: DATA DICTIONARY (6 Pages) ---
    add_heading("7. Database Design & Data Dictionary", level=1, page_break_before=True)
    doc.add_paragraph("The physical design of the NoSQL database involves creating logical 'Collections' mapping to Mongoose Schemas. Below is the exhaustive structural definition (Data Dictionary) of each major collection utilized in the Gurukul Edu WebApp.")

    def add_large_dd_table(name, desc, rows):
        add_heading(f"Collection: {name}", level=2, page_break_before=True)
        doc.add_paragraph(f"Description: {desc}")
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(["Field Name", "Data Type", "Constraint", "Description"]):
            hdr_cells[i].text = col_name
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)

    add_large_dd_table("Users", "Stores authentication and profile data for all roles.", [
        ("_id", "ObjectId", "Primary Key", "Globally unique 12-byte identifier."),
        ("name", "String", "Required, Trim", "User's full legal name."),
        ("email", "String", "Required, Unique", "Used for system login. Must be valid regex format."),
        ("password", "String", "Required, Min: 6", "Bcrypt hashed string (Salt rounds: 10)."),
        ("role", "String", "Enum", "Must be 'admin', 'teacher', or 'student'."),
        ("profilePicture", "String", "Optional", "URL pointing to CDN image file."),
        ("isVerified", "Boolean", "Default: true", "Flag for account suspension/activation."),
        ("createdAt", "Date", "Auto", "Timestamp of document creation.")
    ])

    add_large_dd_table("Batches", "Logical grouping of students and assigned teachers.", [
        ("_id", "ObjectId", "Primary Key", "Unique Batch ID."),
        ("name", "String", "Required, Unique", "E.g., 'Class 10 Science B1'."),
        ("description", "String", "Optional", "Syllabus or timing details."),
        ("teachers", "Array of ObjectId", "Ref: Users", "Teachers authorized to manage this batch."),
        ("students", "Array of ObjectId", "Ref: Users", "Students enrolled in this batch."),
        ("status", "String", "Enum", "'active' or 'inactive'.")
    ])

    add_large_dd_table("Attendance", "Daily attendance tracking records.", [
        ("_id", "ObjectId", "Primary Key", "Unique Record ID."),
        ("studentId", "ObjectId", "Required, Ref: Users", "The student being marked."),
        ("batchId", "ObjectId", "Required, Ref: Batches", "The contextual batch."),
        ("date", "Date", "Required", "Exact date of the lecture/class."),
        ("status", "String", "Enum", "'Present', 'Absent', or 'Late'."),
        ("markedBy", "ObjectId", "Ref: Users", "Admin/Teacher who created the record.")
    ])

    add_large_dd_table("Materials", "Academic resources uploaded by teachers.", [
        ("_id", "ObjectId", "Primary Key", "Unique Resource ID."),
        ("title", "String", "Required", "Title of the note or video."),
        ("fileUrl", "String", "Required", "Secure Cloudinary HTTPS URL string."),
        ("fileType", "String", "Enum", "'PDF', 'Video', 'Note', 'Image'."),
        ("batchId", "ObjectId", "Required, Ref: Batches", "Batch that has access to this file."),
        ("teacherId", "ObjectId", "Required, Ref: Users", "The uploader of the material.")
    ])

    add_large_dd_table("Fees", "Financial records for individual students.", [
        ("_id", "ObjectId", "Primary Key", "Unique Fee Record ID."),
        ("studentId", "ObjectId", "Required, Ref: Users", "The student who owes fees."),
        ("batchId", "ObjectId", "Required, Ref: Batches", "Associated batch fee structure."),
        ("totalAmount", "Number", "Required", "Total course fee (e.g., 50000)."),
        ("amountPaid", "Number", "Default: 0", "Amount paid till date."),
        ("status", "String", "Auto-calculated", "'Paid', 'Pending', or 'Overdue'.")
    ])

    # --- CHAPTER 7: IMPLEMENTATION (3 Pages) ---
    add_heading("8. Implementation & MVC Architecture", level=1, page_break_before=True)
    doc.add_paragraph("The system relies on the Model-View-Controller (MVC) architectural pattern, adapted for RESTful APIs. This ensures a clean separation of concerns, making the codebase highly maintainable and scalable.")
    
    add_heading("Backend Structure (Node.js & Express)", level=2)
    doc.add_paragraph("- Models: Defined using Mongoose. These represent the logical structure of the database collections and include validation logic (e.g., ensuring emails are unique).")
    doc.add_paragraph("- Controllers: Contain the core business logic. When a route is hit, the controller processes the request, interacts with the Model, and returns a JSON response.")
    doc.add_paragraph("- Routes: Map HTTP methods (GET, POST, PUT, DELETE) to specific Controller functions.")
    doc.add_paragraph("- Middleware: Functions that run before controllers. E.g., `authMiddleware.js` intercepts requests, extracts the JWT from headers, verifies it, and attaches the user payload to `req.user`.")

    add_heading("Frontend Structure (React.js)", level=2, page_break_before=True)
    doc.add_paragraph("- Components: Reusable UI elements (e.g., Navbar, Sidebar, Buttons) styled with Tailwind CSS.")
    doc.add_paragraph("- Pages: Higher-level components representing distinct views (e.g., AdminDashboard.jsx, StudentLogin.jsx).")
    doc.add_paragraph("- State Management: Managed using React Hooks (`useState`, `useEffect`, `useContext`) to ensure the UI reacts instantly to data changes without page reloads.")
    doc.add_paragraph("- Axios Interceptors: Automatically attach the JWT token to every outgoing HTTP request header to ensure secure communication with the backend API.")

    # --- CHAPTER 8: TESTING (3 Pages) ---
    add_heading("9. Software Testing & Test Cases", level=1, page_break_before=True)
    doc.add_paragraph("Software testing is critical to ensure the system behaves as expected under various conditions. A rigorous manual testing strategy was employed, focusing on Unit Testing and System Integration Testing.")

    def add_test_table(title, rows):
        add_heading(title, level=3)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = ["TC ID", "Action / Scenario", "Expected Result", "Actual Result", "Status"]
        for i, h in enumerate(hdr):
            table.rows[0].cells[i].text = h
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        doc.add_paragraph()

    add_test_table("Test Cases: Authentication Module", [
        ("TC-01", "Login with valid Email and Password", "System generates JWT and redirects to Dashboard", "JWT generated, successful redirect", "PASS"),
        ("TC-02", "Login with invalid Password", "System returns 401 Unauthorized, shows error toast", "Returned 401, error displayed", "PASS"),
        ("TC-03", "Access protected route without Token", "System denies access, redirects to Login", "Access denied, redirected", "PASS")
    ])

    doc.add_page_break()
    add_test_table("Test Cases: Academic Module", [
        ("TC-04", "Teacher uploads PDF > 10MB", "System rejects file, shows 'Size Limit Exceeded' error", "File rejected correctly", "PASS"),
        ("TC-05", "Teacher marks attendance", "DB updates status for all selected students", "Status updated successfully", "PASS"),
        ("TC-06", "Student attempts to upload material", "System blocks action (RBAC verification fails)", "Action blocked, 403 Forbidden", "PASS")
    ])

    # --- CHAPTER 9: SCREENSHOTS (5 Pages) ---
    add_heading("10. System Screenshots", level=1, page_break_before=True)
    doc.add_paragraph("This section visually demonstrates the developed Graphical User Interface (GUI) of the Gurukul Edu WebApp.")
    
    for i in range(1, 6):
        doc.add_page_break()
        p = doc.add_paragraph(f"\n\n\n\n\n\n<< PLEASE PASTE SCREENSHOT {i} HERE >>\n\n\n\n\n\n")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold, p.runs[0].font.color.rgb = True, RGBColor(41, 128, 185)

    # --- CHAPTER 10: CONCLUSION & BIBLIOGRAPHY (2 Pages) ---
    add_heading("11. Conclusion & Future Enhancements", level=1, page_break_before=True)
    doc.add_paragraph("Conclusion:\nThe 'Gurukul Edu WebApp' successfully fulfills its objective of providing a robust, centralized ERP solution for educational institutes. By leveraging the non-blocking architecture of Node.js and the dynamic rendering capabilities of React.js, the system delivers exceptional performance. The integration of role-based security, cloud media storage, and automated attendance tracking significantly reduces administrative overhead.")
    doc.add_paragraph("Future Scope:\n1. Payment Gateway Integration: Integrating Razorpay or Stripe to allow students to pay fees directly through the portal.\n2. Automated Email/SMS Alerts: Sending real-time notifications to parents if a student is marked absent.\n3. Mobile Application: Developing a cross-platform mobile app using React Native for greater accessibility.\n4. AI Analytics: Implementing Machine Learning models to predict student performance based on historical test data.")

    add_heading("12. Bibliography", level=1, page_break_before=True)
    doc.add_paragraph("The following resources were instrumental in the research and development of this project:")
    refs = [
        "React Official Documentation: https://react.dev/",
        "Node.js API Reference: https://nodejs.org/en/docs/",
        "MongoDB & Mongoose Guidelines: https://mongoosejs.com/docs/guide.html",
        "JSON Web Tokens Standard: https://jwt.io/introduction",
        "Cloudinary Media API: https://cloudinary.com/documentation",
        "Tailwind CSS Layouts: https://tailwindcss.com/docs",
        "Sommerville, I. (2015). 'Software Engineering'. 10th Edition. Pearson."
    ]
    for r in refs:
        doc.add_paragraph(r, style='List Bullet')

    # Save
    doc.save("Gurukul_Project_Documentation_Massive_30Pages.docx")
    print("Successfully generated Gurukul_Project_Documentation_Massive_30Pages.docx!")

if __name__ == "__main__":
    create_massive_document()
