from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_premium_document():
    doc = Document()
    
    # Configure global styling
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Cambria'
    font.size = Pt(11)

    # Helper functions for styling
    def add_heading(text, level, color=(31, 73, 125)):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Calibri Light'
            if color:
                run.font.color.rgb = RGBColor(*color)
            if level == 1:
                run.bold = True
                run.font.size = Pt(18)
        return h

    def add_blank_page(title):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n\n\n\n\n\n[ THIS PAGE IS INTENTIONALLY LEFT BLANK FOR THE PLACEMENT OF: ]\n\n{title}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(127, 140, 141)

    # --- BLANK PAGES ---
    add_blank_page("(1) Project Title Page (colour copy)")
    add_blank_page("(2) College Originality Certificate (colour copy)")
    add_blank_page("(3) Internship Allotment letter (colour copy)")
    add_blank_page("(4) Internship Completion Certificate (colour copy)")

    # --- PREFACE ---
    doc.add_page_break()
    add_heading("Preface", level=1)
    doc.add_paragraph(
        "The evolution of the digital era has revolutionized administrative and academic operations across the globe. "
        "It gives me immense pride to present this comprehensive project documentation on the 'Gurukul Edu WebApp'. "
        "This project represents a crucial intersection between modern software engineering practices and the pressing "
        "need for centralized, efficient educational management systems."
    )
    doc.add_paragraph(
        "Conceptualized as an enterprise-grade solution, this WebApp was designed to automate manual workflows, mitigate "
        "data redundancy, and foster real-time communication between administrators, educators, and students. "
        "Throughout the development lifecycle, rigorous methodologies ranging from requirement gathering to deployment "
        "were adhered to, ensuring a scalable and robust architecture."
    )
    doc.add_paragraph(
        "This document serves as an exhaustive blueprint of the system, detailing its architectural nuances, "
        "data flow logic, backend algorithms, and advanced database structuring."
    )

    # --- ACKNOWLEDGEMENT ---
    doc.add_page_break()
    add_heading("Acknowledgement", level=1)
    doc.add_paragraph(
        "The successful completion of this ambitious project would not have been possible without the invaluable guidance "
        "and unwavering support of several individuals and institutions."
    )
    doc.add_paragraph(
        "I would like to extend my deepest gratitude to the esteemed faculty and management of my college for providing "
        "a conducive academic environment and the theoretical foundation necessary to tackle complex software engineering challenges."
    )
    doc.add_paragraph(
        "Special thanks are owed to my Project Guide for their continuous mentorship. Their constructive criticism, "
        "technical insights into the MERN stack, and strategic direction were the cornerstones of this project's success."
    )
    doc.add_paragraph(
        "I am also profoundly thankful to the development team and my peers during the internship for their collaborative "
        "spirit, code reviews, and for demonstrating industry-standard agile practices."
    )

    # --- EXECUTIVE SUMMARY / INTERNSHIP DOC ---
    doc.add_page_break()
    add_heading("Executive Summary (Internship Experience)", level=1)
    doc.add_paragraph(
        "During the course of my internship, I was entrusted with the end-to-end development of the Gurukul Edu WebApp. "
        "This role transitioned my theoretical understanding into practical, production-ready full-stack development skills."
    )
    add_heading("Technical Acumen Acquired:", level=3)
    doc.add_paragraph("1. Frontend Architecture (React.js): Designed stateful, component-driven user interfaces. Implemented React Router for seamless navigation and utilized Tailwind CSS to ensure a fully responsive, mobile-first design language.", style='List Bullet')
    doc.add_paragraph("2. Backend & API Engineering (Node.js/Express): Architected RESTful APIs with strict validation logic. Mastered the implementation of Middleware for error handling, file parsing (Multer), and route protection.", style='List Bullet')
    doc.add_paragraph("3. Database Modeling (MongoDB/Mongoose): Transitioned from relational thinking to NoSQL document-based modeling. Engineered complex schemas involving references (ObjectIds), indexing for optimization, and Mongoose pre/post save hooks.", style='List Bullet')
    doc.add_paragraph("4. Security Paradigms: Implemented JSON Web Tokens (JWT) for stateless authentication. Secured passwords using bcrypt hashing algorithms and mitigated common web vulnerabilities (CORS, XSS).", style='List Bullet')
    doc.add_paragraph("5. Cloud Integrations: Successfully integrated third-party APIs such as Cloudinary for secure, cloud-based storage of academic materials and multimedia assets.", style='List Bullet')

    # --- INDEX ---
    doc.add_page_break()
    add_heading("Index", level=1)
    index_items = [
        "1. Project Profile & Tech Stack Justification", 
        "2. System Introduction & Objective", 
        "3. Feasibility Study", 
        "4. Existing System vs Proposed System", 
        "5. Hardware & Software Requirements", 
        "6. System Architecture & Flowchart",
        "7. Data Flow Diagram (Levels 0, 1, 2)", 
        "8. Entity Relationship (ER) Diagram", 
        "9. Data Dictionary (Schema Definition)",
        "10. UML Diagrams (Use Case, Class, Sequence)", 
        "11. Module Description (Logic & Integration)", 
        "12. System Security & Authentication",
        "13. Software Testing Strategies",
        "14. System Screenshots", 
        "15. Conclusion & Future Scope", 
        "16. Bibliography & References"
    ]
    for item in index_items:
        doc.add_paragraph(item, style='List Number')

    # --- 1. PROJECT PROFILE ---
    doc.add_page_break()
    add_heading("1. Project Profile & Tech Stack Justification", level=1)
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    records = [
        ("Project Title", "Gurukul Edu WebApp - Enterprise Institution Management"),
        ("Architecture Pattern", "MVC (Model-View-Controller) & RESTful Micro-services"),
        ("Frontend Technology", "React.js (v18+), Tailwind CSS, Vite"),
        ("Backend Technology", "Node.js environment, Express.js framework"),
        ("Database / Storage", "MongoDB Atlas (NoSQL DBaaS), Cloudinary CDN"),
        ("Security Standards", "JWT (Stateless Auth), Bcrypt Hashing, CORS filtering")
    ]
    for i, (k, v) in enumerate(records):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = v

    add_heading("Tech Stack Justification:", level=3)
    doc.add_paragraph("The MERN stack was chosen specifically for its asynchronous, event-driven architecture which is highly suitable for real-time applications like an educational portal. Using JavaScript across the entire stack ensures rapid development, easier code maintenance, and seamless JSON data transmission between client and server.")

    # --- 2. INTRODUCTION ---
    add_heading("2. System Introduction & Objective", level=1)
    doc.add_paragraph(
        "Gurukul Edu WebApp is an integrated, cloud-deployed Educational Resource Planning (ERP) platform. "
        "Its primary objective is to dissolve the communication barriers between the administrative wing, the teaching staff, and the student body."
    )
    doc.add_paragraph(
        "Core Objectives include:\n"
        "• Centralizing user access via Role-Based Access Control (RBAC).\n"
        "• Automating attendance tracking and generating real-time analytics.\n"
        "• Establishing a secure repository for study material distribution.\n"
        "• Digitizing financial tracking to monitor fee submissions and defaults seamlessly."
    )

    # --- 3. FEASIBILITY STUDY ---
    add_heading("3. Feasibility Study", level=1)
    doc.add_paragraph("A thorough feasibility study was conducted prior to the commencement of development to ensure project viability.")
    add_heading("A. Technical Feasibility:", level=3)
    doc.add_paragraph("The proposed system utilizes open-source technologies (Node, React, MongoDB) ensuring high technical feasibility. Cloud deployment on platforms like Render or Vercel requires minimal infrastructure overhead.")
    add_heading("B. Economic Feasibility:", level=3)
    doc.add_paragraph("By eliminating paper-based processes and manual data entry errors, the system significantly reduces operational costs. The use of free-tier cloud services (MongoDB Atlas Free Tier, Cloudinary) ensures minimal initial investment.")
    add_heading("C. Operational Feasibility:", level=3)
    doc.add_paragraph("The UI is engineered to be highly intuitive, adhering to modern UX principles. Thus, administrative staff and teachers require minimal training to transition to the new system.")

    # --- 4. EXISTING VS PROPOSED ---
    add_heading("4. Existing System vs Proposed System", level=1)
    comp_table = doc.add_table(rows=5, cols=3)
    comp_table.style = 'Table Grid'
    comp_headers = ["Parameter", "Existing Manual System", "Proposed WebApp System"]
    for i, header in enumerate(comp_headers):
        comp_table.rows[0].cells[i].text = header
        comp_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        
    comp_data = [
        ("Data Storage", "Decentralized (Files, Registers, Local Excel)", "Centralized Cloud Database (MongoDB Atlas)"),
        ("Accessibility", "Restricted to physical campus premises", "24/7 Global Access via Web Interface"),
        ("Data Security", "High risk of physical loss or unauthorized access", "Encrypted Passwords, Role-based route protection"),
        ("Analytics", "Manual calculation of attendance/fees (Time-consuming)", "Automated dashboard analytics in real-time")
    ]
    for i, (p, e, prop) in enumerate(comp_data, start=1):
        comp_table.rows[i].cells[0].text = p
        comp_table.rows[i].cells[1].text = e
        comp_table.rows[i].cells[2].text = prop

    # --- 5. HW/SW REQ ---
    doc.add_page_break()
    add_heading("5. Hardware & Software Requirements", level=1)
    add_heading("Server-Side (Deployment) Requirements:", level=3)
    doc.add_paragraph("• CPU: 1 vCPU (Minimum)\n• RAM: 512MB for Node runtime\n• OS: Ubuntu Linux 20.04 LTS (or equivalent cloud instance)\n• Database: Minimum 512MB storage cluster (Atlas Free Tier)")
    add_heading("Client-Side Requirements:", level=3)
    doc.add_paragraph("• Device: PC, Laptop, Tablet, or Smartphone\n• Browser: Modern Web Browser supporting ES6+ (Chrome 80+, Safari 13+)\n• Network: Minimum 2 Mbps internet connection for smooth media rendering.")

    # --- DIAGRAM PLACEHOLDERS ---
    def add_diagram_placeholder(title, instruction):
        doc.add_page_break()
        add_heading(title, level=1)
        p = doc.add_paragraph(f"<< PLEASE PASTE THE {instruction} IMAGE HERE FROM diagrams_premium.html >>")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(231, 76, 60)

    add_diagram_placeholder("6. System Architecture & Flowchart", "SYSTEM ARCHITECTURE & FLOWCHART")
    
    doc.add_page_break()
    add_heading("7. Data Flow Diagram (DFD)", level=1)
    p = doc.add_paragraph("<< PLEASE PASTE THE DFD LEVEL 0 (CONTEXT) IMAGE HERE >>\n\n")
    p.runs[0].bold, p.runs[0].font.color.rgb = True, RGBColor(231, 76, 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph("<< PLEASE PASTE THE DFD LEVEL 1 IMAGE HERE >>\n\n")
    p2.runs[0].bold, p2.runs[0].font.color.rgb = True, RGBColor(231, 76, 60)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3 = doc.add_paragraph("<< PLEASE PASTE THE DFD LEVEL 2 IMAGE HERE >>")
    p3.runs[0].bold, p3.runs[0].font.color.rgb = True, RGBColor(231, 76, 60)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_diagram_placeholder("8. Entity Relationship (ER) Diagram", "ER DIAGRAM")

    # --- 9. DATA DICTIONARY ---
    doc.add_page_break()
    add_heading("9. Data Dictionary (Comprehensive Schema Definition)", level=1)
    doc.add_paragraph("The database is normalized and structured utilizing Mongoose ODM. Below are the detailed schema constraints.")

    def add_dd_table(title, rows):
        add_heading(title, level=3)
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        columns = ["Field Name", "Data Type", "Constraint / Default", "Description"]
        for i, col_name in enumerate(columns):
            hdr_cells[i].text = col_name
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        doc.add_paragraph()

    add_dd_table("Table 1: users (Collection)", [
        ("_id", "ObjectId", "PK, Auto-generated", "Globally unique identifier"),
        ("name", "String", "Required, Trimmed", "Full name of user"),
        ("email", "String", "Required, Unique, Lowercase", "Authentication ID"),
        ("password", "String", "Required, Min 6 chars", "Bcrypt hashed cipher"),
        ("role", "String", "Enum: ['admin', 'teacher', 'student']", "Defines RBAC privileges"),
        ("isVerified", "Boolean", "Default: true/false", "Account active status")
    ])

    add_dd_table("Table 2: batches (Collection)", [
        ("_id", "ObjectId", "PK", "Batch unique ID"),
        ("name", "String", "Required, Unique", "Batch nomenclature"),
        ("teachers", "Array[ObjectId]", "Ref: 'users'", "Assigned teaching staff"),
        ("students", "Array[ObjectId]", "Ref: 'users'", "Enrolled students"),
        ("status", "String", "Enum: ['active', 'inactive']", "Current operational state")
    ])

    add_dd_table("Table 3: attendances (Collection)", [
        ("_id", "ObjectId", "PK", "Record ID"),
        ("studentId", "ObjectId", "Required, Ref: 'users'", "Target student"),
        ("batchId", "ObjectId", "Required, Ref: 'batches'", "Contextual batch"),
        ("date", "Date", "Required", "Timestamp of record"),
        ("status", "String", "Enum: ['Present', 'Absent', 'Late']", "Attendance state"),
        ("markedBy", "ObjectId", "Ref: 'users'", "Teacher who logged it")
    ])

    add_dd_table("Table 4: materials (Collection)", [
        ("_id", "ObjectId", "PK", "Resource ID"),
        ("title", "String", "Required", "Document Title"),
        ("fileUrl", "String", "Required", "Cloudinary secure HTTPS URL"),
        ("fileType", "String", "Enum: ['PDF', 'Image', 'Video', 'Note']", "MIME mapping"),
        ("batchId", "ObjectId", "Required, Ref: 'batches'", "Target audience")
    ])

    # --- 10. UML DIAGRAMS ---
    doc.add_page_break()
    add_heading("10. UML Diagrams", level=1)
    p = doc.add_paragraph("<< PLEASE PASTE THE UML CLASS DIAGRAM IMAGE HERE >>\n\n")
    p.runs[0].bold, p.runs[0].font.color.rgb = True, RGBColor(231, 76, 60)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2 = doc.add_paragraph("<< PLEASE PASTE THE UML SEQUENCE DIAGRAM IMAGE HERE >>")
    p2.runs[0].bold, p2.runs[0].font.color.rgb = True, RGBColor(231, 76, 60)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- 11. MODULE DESCRIPTION ---
    doc.add_page_break()
    add_heading("11. Module Description (Logic & Integration)", level=1)
    
    add_heading("A. Administration Module:", level=3)
    doc.add_paragraph("Acts as the system's nucleus. The admin controller handles high-order CRUD operations. "
                      "It orchestrates user onboarding, constructs hierarchical 'Batches', and maps Teachers to Students. "
                      "Furthermore, the admin engine dictates financial parameters via the Fees module, analyzing deficits and surpluses.")
    
    add_heading("B. Teacher / Academic Module:", level=3)
    doc.add_paragraph("Designed for academic execution. The logic intercepts HTTP POST requests containing multipart form data (files), "
                      "pipes them to Cloudinary CDN, and maps the returned secure URL to specific Batches via the Database. "
                      "Teachers also interface with the Attendance API, updating Boolean metrics for hundreds of students simultaneously using bulk-write operations.")

    add_heading("C. Student / Consumer Module:", level=3)
    doc.add_paragraph("A strictly read-only interface (for academic data) ensuring data immutability. Students execute GET requests "
                      "filtered by their JWT decoded payload (studentId). This module serves dynamic dashboards rendering attendance percentages via aggregations and fetching authorized study material links.")

    # --- 12. SECURITY ---
    add_heading("12. System Security & Authentication", level=1)
    doc.add_paragraph("1. Password Cryptography: Plain text passwords are never stored. The system utilizes bcrypt with a salt round of 10 to hash passwords before DB insertion.")
    doc.add_paragraph("2. Stateless Sessions (JWT): Upon successful validation, the server issues an HttpOnly JSON Web Token. Subsequent API requests are intercepted by an `authMiddleware` which decodes the token, verifies the signature using a server-side secret, and checks RBAC permissions before resolving the route.")

    # --- 13. TESTING ---
    add_heading("13. Software Testing Strategies", level=1)
    doc.add_paragraph("Quality assurance was strictly maintained via multi-tier testing methodologies:")
    doc.add_paragraph("1. Unit Testing: Individual API endpoints (e.g., POST /api/auth/login) were tested in isolation using tools like Postman to ensure valid JSON responses and correct HTTP Status Codes (200, 201, 400, 401, 500).")
    doc.add_paragraph("2. Integration Testing: The seamless flow of data between the React Frontend Axios interceptors and the Node backend routes was verified. Evaluated cross-origin resource sharing (CORS) configurations.")
    doc.add_paragraph("3. System Testing: Full application lifecycle tests were conducted. e.g., Admin creates a Teacher -> Teacher uploads Material -> Student logs in and downloads Material.")

    # --- 14. SCREENSHOTS ---
    add_diagram_placeholder("14. System Screenshots", "YOUR APP SCREENSHOTS")

    # --- 15. CONCLUSION ---
    doc.add_page_break()
    add_heading("15. Conclusion & Future Scope", level=1)
    doc.add_paragraph("Conclusion:", style='Heading 3')
    doc.add_paragraph("The Gurukul Edu WebApp successfully achieves its mandate of digitizing educational administration. By utilizing a robust, non-blocking MERN architecture, the system guarantees high throughput and low latency, ultimately enriching the user experience for all institutional stakeholders.")
    add_heading("Future Scope:", level=3)
    doc.add_paragraph("• Implementation of an AI-driven Chatbot to resolve student queries instantly.\n• Integration of Payment Gateways (RazorPay/Stripe) to facilitate automated online fee transactions.\n• Migration to a microservices architecture using Docker containers for infinite horizontal scalability.")

    # --- 16. BIBLIOGRAPHY ---
    doc.add_page_break()
    add_heading("16. Bibliography & References", level=1)
    doc.add_paragraph("1. Official Documentation, React Library: https://reactjs.org/docs/")
    doc.add_paragraph("2. Official Documentation, Node.js: https://nodejs.org/en/docs/")
    doc.add_paragraph("3. Mongoose ODM API Reference: https://mongoosejs.com/")
    doc.add_paragraph("4. Security guidelines, JSON Web Tokens: https://jwt.io/introduction")
    doc.add_paragraph("5. Tailwind CSS Documentation: https://tailwindcss.com/docs")

    # Save
    doc.save("Gurukul_Project_Documentation_Premium.docx")
    print("Successfully generated Gurukul_Project_Documentation_Premium.docx!")

if __name__ == "__main__":
    create_premium_document()
