from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

def create_document():
    doc = Document()
    
    # Configure styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    # Helper functions
    def add_heading(text, level, color=(44, 62, 80)):
        h = doc.add_heading(text, level=level)
        for run in h.runs:
            run.font.name = 'Arial'
            if color:
                run.font.color.rgb = RGBColor(*color)
        return h

    def add_blank_page(title):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"\n\n\n\n\n[ THIS PAGE IS INTENTIONALLY LEFT BLANK FOR: ]\n\n{title}")
        run.bold = True
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(127, 140, 141)

    # --- BLANK PAGES ---
    # Page 1
    add_blank_page("(1) Project Title Page (colour copy)")
    
    # Page 2
    add_blank_page("(2) College Originality Certificate (colour copy)")

    # --- (3) PREFACE ---
    doc.add_page_break()
    add_heading("Preface", level=1)
    doc.add_paragraph(
        "It gives me immense pleasure to present this internship and project report on 'Gurukul Edu WebApp'. "
        "In today's fast-paced digital era, education management requires efficient, reliable, and real-time solutions. "
        "This project was conceptualized to bridge the gap between manual administrative processes and modern technological capabilities."
    )
    doc.add_paragraph(
        "During my internship, I had the invaluable opportunity to work on cutting-edge MERN stack technologies "
        "(MongoDB, Express.js, React.js, Node.js). The experience of transforming a theoretical concept into a fully "
        "functional, deployed web application has been incredibly enriching. This document encompasses the system's "
        "architecture, design methodologies, and the practical implementation details that bring the 'Gurukul Edu WebApp' to life."
    )

    # --- (4) ACKNOWLEDGEMENT ---
    doc.add_page_break()
    add_heading("Acknowledgement", level=1)
    doc.add_paragraph(
        "I would like to express my profound gratitude to everyone who supported and guided me throughout this project and internship."
    )
    doc.add_paragraph(
        "Firstly, I extend my sincere thanks to the Management and Faculty of my esteemed college for providing me with "
        "the academic foundation and the opportunity to undertake this crucial internship."
    )
    doc.add_paragraph(
        "I am deeply indebted to my Project Guide for their continuous encouragement, insightful feedback, and technical "
        "expertise, which were instrumental in the successful completion of this WebApp."
    )
    doc.add_paragraph(
        "I also wish to thank my internship mentors and the development team for exposing me to industry-standard "
        "practices, code reviews, and agile development methodologies."
    )

    # --- BLANK PAGES ---
    # Page 5
    add_blank_page("(5) Internship Allotment letter (colour copy)")
    
    # Page 6
    add_blank_page("(6) Internship Completion Certificate (colour copy)")

    # --- (7) INTERNSHIP DOCUMENT ---
    doc.add_page_break()
    add_heading("Internship Document", level=1)
    add_heading("My Journey as a Full-Stack Developer Intern", level=2)
    doc.add_paragraph(
        "During my internship, I was tasked with developing the 'Gurukul Edu WebApp', a comprehensive platform "
        "designed to digitalize educational operations. Over the course of my tenure, I transitioned from understanding "
        "basic web development concepts to architecting scalable, full-stack applications."
    )
    add_heading("Key Learning Areas:", level=3)
    doc.add_paragraph(
        "1. Frontend Development: I mastered React.js for building dynamic user interfaces. I learned to manage global "
        "state, implement protected routing, and utilize Tailwind CSS for creating responsive and aesthetically pleasing designs.", style='List Bullet'
    )
    doc.add_paragraph(
        "2. Backend Architecture: Working with Node.js and Express.js taught me how to construct robust RESTful APIs. "
        "I implemented secure authentication using JSON Web Tokens (JWT) and bcrypt for password hashing.", style='List Bullet'
    )
    doc.add_paragraph(
        "3. Database Management: I utilized MongoDB and Mongoose ODM to design complex schemas for Users, Batches, "
        "Attendance, and Fees. I learned the importance of normalization and data integrity.", style='List Bullet'
    )
    doc.add_paragraph(
        "This practical exposure significantly honed my problem-solving skills and prepared me for professional software engineering roles."
    )

    # --- (8) INDEX ---
    doc.add_page_break()
    add_heading("Index", level=1)
    index_items = [
        "1) Project Profile", "2) Introduction to system", "3) Existing System", 
        "4) Proposed System", "5) Hardware & Software Requirement", "6) System flowchart",
        "7) Data Flow Diagram (Levels 0, 1, 2)", "8) ER – Diagram", "9) Data Dictionary",
        "10) Module of System (Admin & User)", "11) UML", "12) Coding",
        "13) Screenshots", "14) Reports", "15) Future Enhancement", "16) Bibliography"
    ]
    for item in index_items:
        doc.add_paragraph(item, style='List Number')

    # --- CORE SECTIONS ---
    
    # 1. Project Profile
    doc.add_page_break()
    add_heading("1. Project Profile", level=2)
    p = doc.add_paragraph()
    p.add_run("Project Title: ").bold = True
    p.add_run("Gurukul Edu WebApp\n")
    p.add_run("Technology Stack: ").bold = True
    p.add_run("MERN (MongoDB, Express.js, React.js, Node.js)\n")
    p.add_run("Frontend: ").bold = True
    p.add_run("React.js, Tailwind CSS\n")
    p.add_run("Backend: ").bold = True
    p.add_run("Node.js, Express.js\n")
    p.add_run("Database: ").bold = True
    p.add_run("MongoDB Atlas\n")
    p.add_run("Authentication: ").bold = True
    p.add_run("JWT (JSON Web Tokens)\n")

    # 2. Intro
    add_heading("2. Introduction to system", level=2)
    doc.add_paragraph(
        "Gurukul Edu WebApp is a centralized, cloud-based educational management system. "
        "It provides a seamless interface for Administrators, Teachers, and Students to interact, "
        "manage academic records, track attendance, and process fees digitally."
    )

    # 3. Existing
    add_heading("3. Existing System", level=2)
    doc.add_paragraph(
        "The current system in many institutes relies heavily on manual paperwork and decentralized Excel sheets. "
        "This leads to data redundancy, risk of document loss, delayed communication between teachers and students, "
        "and tedious fee tracking processes."
    )

    # 4. Proposed
    add_heading("4. Proposed System", level=2)
    doc.add_paragraph(
        "The proposed Gurukul Edu WebApp automates these processes by providing a centralized database. "
        "Features include automated fee status updates, instant study material sharing, real-time attendance tracking, "
        "and secure role-based access control, thereby increasing overall institutional efficiency."
    )

    # 5. HW/SW Req
    add_heading("5. Hardware & Software Requirement", level=2)
    doc.add_paragraph("Hardware Requirements:", style='Heading 3')
    doc.add_paragraph("Processor: Intel Core i3 or equivalent (Minimum)\nRAM: 4 GB (8 GB Recommended)\nStorage: 50 GB Free Space", style='List Bullet')
    doc.add_paragraph("Software Requirements:", style='Heading 3')
    doc.add_paragraph("OS: Windows 10/11, macOS, or Linux\nBrowser: Google Chrome, Firefox, or Safari\nEnvironment: Node.js (v18+)\nDatabase: MongoDB Atlas", style='List Bullet')

    # 6. Flowchart
    doc.add_page_break()
    add_heading("6. System Flowchart", level=2)
    p = doc.add_paragraph("<< PLEASE PASTE THE SYSTEM FLOWCHART IMAGE HERE FROM diagrams_perfect.html >>")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    # 7. DFD
    doc.add_page_break()
    add_heading("7. Data Flow Diagram", level=2)
    p = doc.add_paragraph("<< PLEASE PASTE THE DFD LEVEL 0 (CONTEXT) IMAGE HERE FROM diagrams_perfect.html >>\n\n\n")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    p2 = doc.add_paragraph("<< PLEASE PASTE THE DFD LEVEL 1 IMAGE HERE FROM diagrams_perfect.html >>\n\n\n")
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True
    p2.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    p3 = doc.add_paragraph("<< PLEASE PASTE THE DFD LEVEL 2 IMAGE HERE FROM diagrams_perfect.html >>")
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.runs[0].bold = True
    p3.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    # 8. ERD
    doc.add_page_break()
    add_heading("8. ER – Diagram", level=2)
    p = doc.add_paragraph("<< PLEASE PASTE THE ER DIAGRAM IMAGE HERE FROM diagrams_perfect.html >>")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    # 9. Data Dictionary
    doc.add_page_break()
    add_heading("9. Data Dictionary", level=2)
    doc.add_paragraph("The database consists of the following primary collections:")

    def add_dd_table(title, columns, rows):
        add_heading(title, level=3)
        table = doc.add_table(rows=1, cols=len(columns))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        for i, col_name in enumerate(columns):
            hdr_cells[i].text = col_name
            hdr_cells[i].paragraphs[0].runs[0].bold = True
        
        for row_data in rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        doc.add_paragraph()

    add_dd_table("1. Users Table (Collection)", 
                 ["Field Name", "Data Type", "Constraint", "Description"],
                 [
                     ("_id", "ObjectId", "Primary Key", "Unique identifier for user"),
                     ("name", "String", "Required", "Full name of the user"),
                     ("email", "String", "Unique, Required", "Login email"),
                     ("password", "String", "Required", "Hashed password"),
                     ("role", "String", "Enum", "admin, teacher, or student")
                 ])

    add_dd_table("2. Batches Table (Collection)", 
                 ["Field Name", "Data Type", "Constraint", "Description"],
                 [
                     ("_id", "ObjectId", "Primary Key", "Unique identifier for batch"),
                     ("name", "String", "Required", "Name of the batch (e.g. Science 2024)"),
                     ("teachers", "Array[ObjectId]", "Foreign Key", "References to Users"),
                     ("status", "String", "Enum", "active or inactive")
                 ])

    add_dd_table("3. Attendance Table (Collection)", 
                 ["Field Name", "Data Type", "Constraint", "Description"],
                 [
                     ("_id", "ObjectId", "Primary Key", "Unique identifier"),
                     ("studentId", "ObjectId", "Foreign Key", "References Student"),
                     ("batchId", "ObjectId", "Foreign Key", "References Batch"),
                     ("date", "Date", "Required", "Date of attendance"),
                     ("status", "String", "Enum", "Present, Absent, Late")
                 ])

    # 10. Modules
    doc.add_page_break()
    add_heading("10. Module of System (Admin & User)", level=2)
    doc.add_paragraph("Admin Module:", style='Heading 3')
    doc.add_paragraph("- Dashboard: Overall statistics.\n- User Management: Add/Edit Teachers and Students.\n- Batch Management: Create batches and assign members.\n- Fee Management: Assign and track fee payments.", style='List Bullet')
    
    doc.add_paragraph("Teacher Module:", style='Heading 3')
    doc.add_paragraph("- Academic Hub: Mark and track daily attendance.\n- Material Upload: Share PDFs, Notes, and Videos with assigned batches.\n- Evaluation: Conduct tests and upload results.", style='List Bullet')

    doc.add_paragraph("Student Module:", style='Heading 3')
    doc.add_paragraph("- Profile: View personal details.\n- Resources: Download study materials.\n- Trackers: View personal attendance, test scores, and fee status.", style='List Bullet')

    # 11. UML
    doc.add_page_break()
    add_heading("11. UML", level=2)
    p = doc.add_paragraph("<< PLEASE PASTE THE UML DIAGRAM IMAGE HERE FROM diagrams_perfect.html >>")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    # 12. Coding
    doc.add_page_break()
    add_heading("12. Coding", level=2)
    doc.add_paragraph("Sample Code Snippet (Mongoose Schema):")
    code = doc.add_paragraph(
        "const userSchema = new mongoose.Schema({\n"
        "  name: { type: String, required: true },\n"
        "  email: { type: String, required: true, unique: true },\n"
        "  password: { type: String, required: true },\n"
        "  role: { type: String, enum: ['admin', 'teacher', 'student'] }\n"
        "});"
    )
    for run in code.runs:
        run.font.name = 'Courier New'

    # 13. Screenshots
    doc.add_page_break()
    add_heading("13. Screenshots", level=2)
    p = doc.add_paragraph("<< PLEASE PASTE YOUR ACTUAL APP SCREENSHOTS HERE >>")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = RGBColor(231, 76, 60)

    # 14. Reports
    doc.add_page_break()
    add_heading("14. Reports", level=2)
    doc.add_paragraph(
        "The system generates various analytical and tabular reports including:\n"
        "1. Monthly Attendance Reports for specific batches.\n"
        "2. Defaulter Lists for students with pending fees.\n"
        "3. Performance Reports based on test results."
    )

    # 15. Future Enhancement
    add_heading("15. Future Enhancement", level=2)
    doc.add_paragraph(
        "- Integration of Payment Gateways (Stripe/Razorpay) for direct fee payments.\n"
        "- Mobile Application (React Native) for students.\n"
        "- AI-based performance prediction algorithms."
    )

    # 16. Bibliography
    add_heading("16. Bibliography", level=2)
    doc.add_paragraph(
        "1. React Documentation: https://reactjs.org/\n"
        "2. Node.js Official Docs: https://nodejs.org/\n"
        "3. MongoDB Atlas Guide: https://docs.atlas.mongodb.com/\n"
        "4. Tailwind CSS Typography: https://tailwindcss.com/"
    )

    # Save
    doc.save("Gurukul_Project_Documentation.docx")
    print("Successfully generated Gurukul_Project_Documentation.docx!")

if __name__ == "__main__":
    create_document()
